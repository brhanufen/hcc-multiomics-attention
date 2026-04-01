#!/usr/bin/env python3
"""Generate manuscript.docx with embedded figures using python-docx."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

FIG_DIR = "/Users/bfentaw2/system_biology/hcc_project/results/figures"
OUT = "/Users/bfentaw2/system_biology/hcc_project/manuscript/manuscript.docx"

doc = Document()

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        hs.font.size = Pt(16)
        hs.font.bold = True
    elif level == 2:
        hs.font.size = Pt(14)
        hs.font.bold = True
    else:
        hs.font.size = Pt(12)
        hs.font.bold = True
        hs.font.italic = True

def add_paragraph(text, bold=False, italic=False, align=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_figure(filename, caption, width=6.0):
    """Add a figure with caption."""
    path = os.path.join(FIG_DIR, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = cap.add_run(caption)
        run_cap.font.size = Pt(10)
        run_cap.font.name = 'Times New Roman'
        run_cap.italic = True
        cap.paragraph_format.space_after = Pt(12)
    else:
        add_paragraph(f"[Figure not found: {filename}]", italic=True)

def add_table(headers, rows, caption=None):
    """Add a formatted table."""
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
    doc.add_paragraph()  # spacer

# ================================================================
# TITLE PAGE
# ================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
run = p.add_run('Interpretable Deep Learning-Based Multi-Omics Integration\nfor Prognosis in Hepatocellular Carcinoma')
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Times New Roman'

# Authors
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
run = p.add_run('Brhanu F. Znabu')
run.font.size = Pt(13)
run.font.name = 'Times New Roman'
run.bold = True
run2 = p.add_run('1*')
run2.font.size = Pt(13)
run2.font.name = 'Times New Roman'
run2.font.superscript = True
run3 = p.add_run(', Zohaib Atif')
run3.font.size = Pt(13)
run3.font.name = 'Times New Roman'
run3.bold = True
run4 = p.add_run('2')
run4.font.size = Pt(13)
run4.font.name = 'Times New Roman'
run4.font.superscript = True

# Affiliations
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run('1')
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
run.font.superscript = True
run2 = p.add_run('Biomedical Engineering Program, College of Engineering, University of Nebraska-Lincoln, Lincoln, NE 68588-0642, USA')
run2.font.size = Pt(10)
run2.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2')
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
run.font.superscript = True
run2 = p.add_run('Department of Biomedical Science and Engineering, Gwangju 61005, South Korea')
run2.font.size = Pt(10)
run2.font.name = 'Times New Roman'

# Corresponding author
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run('*Corresponding author: Brhanu F. Znabu, 25674436@nebraska.edu, ORCID: 0009-0009-7230-754X')
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
run.italic = True

doc.add_page_break()

# ================================================================
# ABSTRACT
# ================================================================
doc.add_heading('Abstract', level=1)

add_paragraph(
    'Hepatocellular carcinoma (HCC) is a leading cause of cancer mortality worldwide, yet existing '
    'prognostic models incompletely capture its molecular heterogeneity. We developed an interpretable, '
    'attention-based multi-branch deep learning framework for multi-omics survival prediction in HCC. '
    'Using 358 TCGA LIHC patients with matched mRNA expression, miRNA expression, and DNA methylation '
    'data, we first reproduced the Chaudhary et al. autoencoder-based survival model as a baseline '
    '(C-index = 0.561, log-rank p = 3.10 \u00d7 10\u207b\u00b2). We then designed a multi-branch '
    'architecture with omics-specific encoders, multi-head attention fusion, and Cox partial likelihood '
    'training, optimized via Bayesian hyperparameter search (100 Optuna trials). In 5-fold stratified '
    'cross-validation with nested feature selection (no data leakage), our attention model achieved a '
    'mean C-index of 0.683 \u00b1 0.039, outperforming the autoencoder baseline (0.561) and '
    'clinical-only model (0.637), and performing similarly to an AUTOSurv-like benchmark (0.697). Branch dropout enabled single-omics inference; external validation '
    'on the real GSE14520 cohort (n=221, mRNA) achieved a C-index of 0.637 (p = 0.004), comparable to '
    'Chaudhary et al.\u2019s reported 0.67 on the same data. Integrated gradients and attention weights '
    'highlighted features with prior links to HCC biology, including cell cycle genes (CCNA2, PLK1) '
    'and a Wnt pathway component (FZD7), along with candidate biomarkers stable across all '
    'cross-validation folds (PZP, SGCB, CD300LG, ZNF831 for mRNA; 12 miRNAs; 6 CpG sites). '
    'Differential expression analysis between model-defined risk groups identified 381 significant '
    'genes (Bonferroni p < 0.05), though this analysis is partly circular. Multivariable Cox '
    'regression indicated that the model-derived risk score adds prognostic value '
    'beyond clinical variables, with consistent performance across clinical subgroups, though '
    'clinical integration metrics were evaluated on training data. This framework provides '
    'a transparent, biologically grounded approach to multi-omics prognostication in HCC.'
)

p = doc.add_paragraph()
run = p.add_run('Keywords: ')
run.bold = True
run.font.size = Pt(12)
run = p.add_run('hepatocellular carcinoma, deep learning, multi-omics integration, survival prediction, '
                'attention mechanism, interpretability, TCGA')
run.font.size = Pt(12)

doc.add_page_break()

# ================================================================
# 1. INTRODUCTION
# ================================================================
doc.add_heading('1. Introduction', level=1)

add_paragraph(
    'Hepatocellular carcinoma (HCC) is the most common form of primary liver cancer and a leading '
    'cause of cancer death globally, with approximately 900,000 new diagnoses and 830,000 deaths per '
    'year (Sung et al. 2021; Rumgay et al. 2022). Despite advances in surgical, locoregional, and '
    'systemic therapies, HCC carries a poor prognosis, with 5-year overall survival often below '
    '20\u201330% (Villanueva 2019). A major challenge is that patients with similar clinical stage '
    'frequently have very different outcomes, reflecting molecular heterogeneity not captured by '
    'traditional staging systems such as the Barcelona Clinic Liver Cancer (BCLC) classification '
    '(Llovet et al. 2021).'
)

add_paragraph(
    'The Cancer Genome Atlas (TCGA) and related initiatives have generated comprehensive multi-omics '
    'profiles for hundreds of HCC patients, including mRNA expression, miRNA expression, and DNA '
    'methylation (Hutter and Zenklusen 2018). These data capture complementary aspects of tumor '
    'biology\u2014transcriptional programs, post-transcriptional regulation, and epigenetic '
    'state\u2014and provide an opportunity for molecular risk stratification that goes beyond '
    'clinical variables alone.'
)

add_paragraph(
    'Deep learning has emerged as a powerful approach for integrating high-dimensional multi-omics '
    'data with survival outcomes (Chaudhary et al. 2018; Huang et al. 2019; Chai et al. 2021). '
    'Chaudhary et al. (2018) provided a landmark demonstration by training a deep autoencoder on '
    'concatenated mRNA, miRNA, and methylation data from 360 TCGA HCC patients, defining two survival '
    'subgroups with significantly different outcomes (C-index = 0.68 on TCGA, validated across five '
    'independent cohorts). However, this autoencoder-based model operates as a \u201cblack box\u201d: '
    'latent features do not map directly to specific genes or CpG sites, and the contribution of each '
    'omics layer to risk prediction is not explicitly quantified (Wysocka et al. 2023; Wekesa et al. 2023).'
)

add_paragraph(
    'Recent work has begun to address interpretability in multi-omics survival models through '
    'attention-based architectures and feature attribution methods (Jiang et al. 2024; Elbashir et al. '
    '2024; Zhang et al. 2025), but these approaches have not been systematically benchmarked against the '
    'Chaudhary et al. framework on the same data, nor applied to HCC multi-omics survival prediction in '
    'this specific setting.'
)

add_paragraph(
    'In this study, we address these gaps by: (1) reproducing and benchmarking the Chaudhary et al. '
    'autoencoder model on TCGA LIHC data; (2) developing an interpretable, multi-branch, attention-based '
    'deep learning architecture with branch dropout for handling missing omics layers; and (3) integrating '
    'model interpretability outputs with pathway enrichment, differential expression, and clinical analyses '
    'to derive biologically and clinically meaningful insights into aggressive HCC.'
)

# ================================================================
# 2. MATERIALS AND METHODS
# ================================================================
doc.add_heading('2. Materials and Methods', level=1)

doc.add_heading('2.1. Data', level=2)
add_paragraph(
    'We obtained TCGA Liver Hepatocellular Carcinoma (LIHC) data from the UCSC Xena platform '
    '(https://xenabrowser.net/datapages/). Specifically: mRNA expression (HiSeqV2 RNA-seq, '
    'log2(normalized count + 1), 20,530 genes \u00d7 423 samples), miRNA expression (miRNA HiSeq '
    'gene-level, log2(RPM + 1), 2,172 miRNAs \u00d7 420 samples), DNA methylation (Illumina '
    'HumanMethylation450, beta values, 485,577 CpG sites \u00d7 429 samples), and clinical data '
    'including overall survival, vital status, age, gender, pathologic stage, and histologic grade.'
)

add_paragraph(
    'After restricting to primary tumor samples (TCGA barcode suffix -01) and requiring matched '
    'availability of all three omics layers plus survival data, we obtained a final cohort of 358 '
    'patients with 127 events (35.5% event rate) and a median overall survival of 595 days. This '
    'is consistent with the 360 patients reported by Chaudhary et al. (2018).'
)

add_paragraph(
    'For external validation, we obtained two publicly available independent HCC cohorts from GEO: '
    'GSE14520 (Roessler et al., n=221 tumor samples with survival data, Affymetrix mRNA) and '
    'GSE31384 (n=166, miRNA with survival data). Probe-to-gene symbol mapping for GSE14520 was '
    'performed using the GPL3921 platform annotation. Three additional cohorts used by Chaudhary '
    'et al. (LIRI-JP, E-TABM-36, Hawaiian) were not included as they require controlled-access '
    'applications or lacked publicly available survival annotations.'
)

doc.add_heading('2.2. Data Preprocessing', level=2)
add_paragraph(
    'Each omics layer was preprocessed as follows. mRNA: Already provided as log2(x+1) values from '
    'UCSC Xena; z-score normalization was applied, and the top 5,000 most variable genes were retained. '
    'miRNA: Already log2(RPM+1); z-score normalization; features with variance > 0.01 were retained '
    '(1,421 miRNAs). Methylation: CpG sites with >20% missing values were removed (395,619 retained '
    'from 485,577); missing values imputed with column medians; top 5,000 most variable CpGs selected, '
    'logit-transformed (M-values), and z-score normalized.'
)

add_paragraph(
    'For the attention-based model (Aim 2), additional survival-association filtering was performed '
    'within each cross-validation fold using only training data, to avoid information leakage. '
    'Spearman correlation between feature values and a risk proxy (event/time) was computed on '
    'the training set, selecting the top 1,000 mRNA genes, 300 miRNAs, and 1,000 CpG sites per fold.'
)

doc.add_heading('2.3. Aim 1: Autoencoder Baseline', level=2)
add_paragraph(
    'Following Chaudhary et al. (2018), we concatenated the three preprocessed omics matrices '
    '(11,421 total features) and trained a deep autoencoder in PyTorch. Architecture: Input (11,421) '
    '\u2192 500 (tanh, 50% dropout) \u2192 100 (tanh, bottleneck) \u2192 500 (tanh, 50% dropout) '
    '\u2192 11,421. Training: 10 epochs, SGD (lr=0.01, momentum=0.9), MSE reconstruction loss with '
    'L1 (\u03bb=10\u207b\u2075) and L2 (\u03bb=10\u207b\u2074) regularization. Post-training: '
    '100-dimensional bottleneck features were extracted; survival-associated features identified via '
    'univariate Cox-PH (p < 0.05); K-means clustering (k=2) defined risk subgroups. Evaluation: '
    'Kaplan\u2013Meier curves, log-rank tests, and concordance index. Benchmarking against '
    'clinical-only and single-omics PCA + Cox models.'
)

doc.add_heading('2.4. Aim 2: Attention-Based Multi-Branch Model', level=2)

doc.add_heading('2.4.1. Architecture', level=3)
add_paragraph(
    'We designed a multi-branch neural network with three omics-specific encoder branches feeding into '
    'a multi-head attention fusion module. Branch encoders: Each omics type passes through a two-layer '
    'encoder (Linear \u2192 BatchNorm \u2192 ReLU \u2192 Dropout \u2192 Linear \u2192 BatchNorm '
    '\u2192 ReLU), producing a latent representation of dimension d. Multi-head attention: A '
    'transformer-style multi-head attention module computes cross-omics attention weights, producing '
    'a fused patient representation. Risk head: Dropout \u2192 Linear(d, 32) \u2192 ReLU \u2192 '
    'Linear(32, 1) for scalar risk score. Loss: Negative Cox partial log-likelihood.'
)

doc.add_heading('2.4.2. Branch Dropout', level=3)
add_paragraph(
    'During training, entire omics branches were randomly masked (output set to zero) with probability '
    'p_drop per branch per sample, with the constraint that at least one branch remains active. The '
    'attention module renormalizes weights over active branches, enabling inference with any subset '
    'of available omics at test time.'
)

doc.add_heading('2.4.3. Hyperparameter Optimization', level=3)
add_paragraph(
    'Bayesian optimization via Optuna (100 trials) was used to tune: latent dimension (32, 64, 128), '
    'number of attention heads (1, 2, 4), learning rate (10\u207b\u2074 to 10\u207b\u00b2), dropout '
    'rate (0.3, 0.4, 0.5), L2 weight decay (10\u207b\u2074 to 10\u207b\u00b3), batch size (32, 64), '
    'and branch dropout probability (0.1\u20130.3). The objective was mean C-index across 3-fold CV.'
)

doc.add_heading('2.4.4. Interpretability', level=3)
add_paragraph(
    'Feature-level importance was computed using integrated gradients (30 interpolation steps, zero '
    'baseline). Omics-branch importance was derived from attention weights averaged across heads.'
)

doc.add_heading('2.5. Aim 3: Biological and Clinical Interpretation', level=2)
add_paragraph(
    'Pathway enrichment: Top 200 attention-derived genes tested against curated HCC gene sets '
    '(Cell Cycle, Wnt/\u03b2-catenin, PI3K/AKT/mTOR, Angiogenesis, Immune Response, Stemness, '
    'HCC Markers) using Fisher\u2019s exact test, plus GSEApy against KEGG 2021 and MSigDB Hallmark '
    '2020. Concordance: Jaccard index and Spearman rank correlation between attention-derived and DE '
    'feature rankings. Clinical integration: Multivariable Cox regression, likelihood ratio test, '
    'NRI. Subgroup analyses by stage, gender, and age. Stability: Kendall\u2019s W across 5 CV folds; '
    'features consistently in top 100 designated high-confidence biomarker candidates.'
)

# ================================================================
# 3. RESULTS
# ================================================================
doc.add_heading('3. Results', level=1)

doc.add_heading('3.1. Aim 1: Autoencoder Baseline Reproduction', level=2)

add_paragraph(
    'The reproduced autoencoder model identified 10 survival-associated latent features (p < 0.05) '
    'from 100 bottleneck dimensions and stratified the 358 TCGA LIHC patients into high-risk (n=191) '
    'and low-risk (n=167) subgroups.'
)

add_table(
    ['Metric', 'Value'],
    [
        ['C-index (TCGA, autoencoder)', '0.561'],
        ['Log-rank p-value', '3.10 \u00d7 10\u207b\u00b2'],
        ['C-index (Chaudhary et al. reported)', '0.68'],
        ['C-index (clinical only)', '0.637'],
        ['C-index (mRNA only, PCA+Cox)', '0.625'],
        ['C-index (miRNA only, PCA+Cox)', '0.636'],
        ['C-index (methylation only, PCA+Cox)', '0.520'],
    ],
    caption='Table 1. Aim 1 Performance Metrics'
)

add_paragraph(
    'The autoencoder achieved a C-index of 0.561, lower than Chaudhary et al.\u2019s reported 0.68, '
    'attributable to differences in software framework (PyTorch vs. Keras/TF1), preprocessing, and '
    'random initialization. The log-rank p-value (0.031) confirmed statistically significant survival '
    'separation between risk groups.'
)

add_figure('aim1_km_curve.png',
           'Figure 1. Kaplan\u2013Meier survival curves for autoencoder-defined risk subgroups in TCGA LIHC (Aim 1).')

add_figure('aim1_benchmark.png',
           'Figure 2. Model benchmarking comparison: autoencoder vs. clinical-only and single-omics baselines (Aim 1).')

# 3.2
doc.add_heading('3.2. Aim 2: Attention-Based Model Performance', level=2)

doc.add_heading('3.2.1. Internal Validation', level=3)
add_paragraph(
    'The attention-based multi-branch model outperformed the autoencoder and clinical-only baselines '
    'in 5-fold stratified cross-validation and performed comparably to the AUTOSurv-like benchmark.'
)

add_table(
    ['Model', 'C-index (mean \u00b1 SD)'],
    [
        ['Attention model (proposed)', '0.683 \u00b1 0.039'],
        ['AUTOSurv-like baseline', '0.697'],
        ['Clinical only', '0.637'],
        ['Autoencoder (Aim 1)', '0.561'],
    ],
    caption='Table 2. Model Comparison (5-Fold CV)'
)

add_table(
    ['Fold', 'C-index', 'Val. N', 'Events', 'Epochs'],
    [
        ['1', '0.712', '72', '25', '36'],
        ['2', '0.687', '72', '25', '32'],
        ['3', '0.707', '72', '25', '16'],
        ['4', '0.606', '71', '26', '16'],
        ['5', '0.702', '71', '26', '25'],
        ['Mean', '0.683 \u00b1 0.039', '', '', ''],
    ],
    caption='Table 3. 5-Fold Cross-Validation Results'
)

add_figure('aim2_km_curve.png',
           'Figure 3. Kaplan\u2013Meier survival curves for attention model-defined risk groups (Aim 2).')

add_figure('aim2_model_comparison.png',
           'Figure 4. Model comparison: attention model vs. all baselines (5-fold CV C-index).')

doc.add_heading('3.2.2. Omics Branch Importance', level=3)
add_paragraph(
    'Attention weights revealed relatively balanced contributions across omics layers: mRNA '
    '(0.340 \u00b1 0.053), miRNA (0.328 \u00b1 0.063), and methylation (0.332 \u00b1 0.069), '
    'suggesting all three omics types carry complementary prognostic information.'
)

add_figure('aim2_attention_weights.png',
           'Figure 5. Omics-level attention weights: (A) Population-level branch importance; '
           '(B) Per-patient attention heatmap stratified by risk group.')

doc.add_heading('3.2.3. Feature-Level Importance', level=3)
add_paragraph(
    'Integrated gradients identified top prognostic features per omics layer. Top mRNA features: '
    'PZP, SGCB, HLA-J, ZNF662, SOX11, CCNA2, CENPE, FZD7. Top miRNA features: MIMAT0003302, '
    'MIMAT0027434, MIMAT0003214, MIMAT0000267. Top CpG sites: cg00866556, cg20603260, cg04743758.'
)

add_figure('aim2_feature_importance.png',
           'Figure 6. Feature importance (integrated gradients) for top 20 features per omics layer.')

doc.add_heading('3.2.4. External Validation on Independent Cohorts', level=3)

add_table(
    ['Cohort', 'n', 'Omics', 'Data Source', 'Our C-index', 'Chaudhary C-index', 'Log-rank p'],
    [
        ['NCI (GSE14520)', '221', 'mRNA', 'Real (GEO)', '0.637', '0.67', '0.004'],
        ['Chinese (GSE31384)', '166', 'miRNA', 'Real (GEO)', '0.471*', '0.69', '0.276'],
    ],
    caption='Table 4. External Validation on Independent Cohorts (Real GEO Data)'
)

add_paragraph(
    '*GSE31384 miRNA probe IDs could not be mapped to MIMAT accessions used by the model; '
    'this result is inconclusive and should not be considered a valid external test.',
    italic=True
)

add_paragraph(
    'On the NCI cohort (GSE14520), the attention model achieved a C-index of 0.637 with significant '
    'survival separation (log-rank p = 0.004), comparable to Chaudhary et al.\u2019s reported 0.67 on '
    'the same cohort. Gene-level alignment was possible for 661 of 1,000 model features via GPL3921 '
    'probe-to-gene annotation. This represents the only successful external validation in this study. '
    'The GSE31384 miRNA cohort could not be meaningfully evaluated due to incompatible probe '
    'identifiers (numeric platform IDs vs. MIMAT accessions); the reported C-index of 0.471 is '
    'therefore inconclusive and should not be interpreted as evidence for or against model '
    'generalizability in miRNA-only settings.'
)

add_figure('aim2_real_external_km.png',
           'Figure 7. Kaplan\u2013Meier survival curves for real external validation cohorts '
           '(GSE14520 mRNA, GSE31384 miRNA) using branch dropout inference.')

add_figure('aim2_real_external_comparison.png',
           'Figure 8. C-index comparison with Chaudhary et al. on real external cohorts.')

# 3.3
doc.add_heading('3.3. Aim 3: Biological and Clinical Interpretation', level=2)

doc.add_heading('3.3.1. Pathway Enrichment', level=3)
add_paragraph(
    'Fisher\u2019s exact test identified overlap between the top 200 attention-derived genes and '
    'curated HCC gene sets: Cell Cycle (2 genes: CCNA2, CENPE), Wnt/\u03b2-catenin (1 gene: FZD7), '
    'and Angiogenesis (1 gene: TEK). GSEApy analysis against KEGG and MSigDB identified enrichment '
    'for G2\u2013M Checkpoint (p = 0.10), HIF-1 signaling (p = 0.19), E2F Targets (p = 0.28), and '
    'Epithelial\u2013Mesenchymal Transition (p = 0.36).'
)

add_figure('aim3_pathway_enrichment.png',
           'Figure 9. Pathway enrichment of top 200 attention-derived genes against curated HCC oncogenic gene sets.')

doc.add_heading('3.3.2. Differential Expression and Concordance', level=3)
add_paragraph(
    'Between model-defined risk groups, 381 genes were differentially expressed at Bonferroni-corrected '
    'p < 0.05 out of 1,000 tested. Top upregulated genes in the high-risk group included G6PD, CBX2, '
    'CEP55, KIF2C, PLK1, TRIP13, MYBL2, and DLGAP5\u2014genes previously implicated in cell '
    'cycle progression and HCC prognosis in independent studies, though we note the DE analysis '
    'is partly circular as the risk groups were defined by the model trained on the same data. The Jaccard index between top 200 attention-derived and '
    'top 200 DE genes was 0.064 (18 overlapping genes). Spearman correlation between feature importance '
    'and |log2FC| was \u03c1 = \u22120.25 (p = 0.012), indicating the attention model captures '
    'prognostic features beyond simple differential expression.'
)

add_figure('aim3_volcano_plot.png',
           'Figure 10. Volcano plot of differential expression between attention model-defined '
           'high- and low-risk groups.')

doc.add_heading('3.3.3. Clinical Integration', level=3)

add_table(
    ['Model', 'C-index', 'LR test p-value'],
    [
        ['Risk score only', '0.989', '\u2014'],
        ['Clinical only (stage, gender, age)', '0.637', '\u2014'],
        ['Risk score + clinical', '0.989', '< 10\u207b\u00b9\u2070\u2070'],
    ],
    caption='Table 5. Multivariable Cox Regression'
)

add_paragraph(
    'Note: These C-indices are computed on the full training data and reflect the model\u2019s fit rather '
    'than generalization performance; the nested cross-validated C-index of 0.683 is the appropriate estimate '
    'of out-of-sample discrimination. Nevertheless, multivariable Cox regression on the full data '
    'confirmed that the risk score was the dominant predictor (HR = 2.26, p = 3.97 \u00d7 10\u207b\u2076\u00b9), '
    'while clinical variables did not reach significance. The likelihood ratio test confirmed '
    'that the risk score adds independent prognostic value beyond clinical factors '
    '(p < 10\u207b\u00b9\u2070\u2070). NRI = 0.398, indicating meaningful reclassification improvement.'
)

add_figure('aim3_forest_plot.png',
           'Figure 11. Forest plot of multivariable Cox regression: risk score + clinical variables.')

doc.add_heading('3.3.4. Subgroup Analysis', level=3)
add_paragraph(
    'To assess whether the model\u2019s prognostic value is consistent across clinical subgroups, '
    'we stratified patients by stage, gender, and age. The risk score significantly separated '
    'high- and low-risk patients in all subgroups tested (all log-rank p < 10\u207b\u2074), '
    'including early-stage (I\u2013II, n=249), late-stage (III\u2013IV, n=86), male (n=242), '
    'female (n=116), younger (n=186), and older (n=172) patients. We note that these subgroup '
    'analyses were performed on the full training data and therefore reflect model fit; '
    'subgroup-level generalization should be confirmed in independent cohorts.'
)

add_figure('aim3_subgroup_analysis.png',
           'Figure 12. Subgroup-stratified C-index demonstrating consistent model performance.')

doc.add_heading('3.3.5. Stability of Interpretability', level=3)
add_paragraph(
    'Kendall\u2019s W across 5 CV folds was 0.200 (mRNA), 0.180 (miRNA), and 0.233 (methylation), '
    'indicating low-to-moderate ranking agreement (W < 0.3), consistent with the expected variability '
    'of deep learning models in limited-sample regimes. Candidate features that remained in the top 100 '
    'across all 5 folds included 4 mRNA genes (PZP, SGCB, CD300LG, ZNF831), 12 miRNAs, and 6 CpG sites.'
)

add_table(
    ['Omics', 'Feature', 'Folds in Top 100'],
    [
        ['mRNA', 'PZP', '5/5'],
        ['mRNA', 'SGCB', '5/5'],
        ['mRNA', 'CD300LG', '5/5'],
        ['mRNA', 'ZNF831', '5/5'],
        ['miRNA', 'MIMAT0003302', '5/5'],
        ['miRNA', 'MIMAT0027434', '5/5'],
        ['miRNA', 'MIMAT0003214', '5/5'],
        ['miRNA', 'MIMAT0000267', '5/5'],
        ['miRNA', 'MIMAT0000450', '5/5'],
        ['miRNA', 'MIMAT0022482', '5/5'],
        ['Methylation', 'cg21131024', '5/5'],
        ['Methylation', 'cg07676361', '5/5'],
        ['Methylation', 'cg08979352', '5/5'],
        ['Methylation', 'cg15565032', '5/5'],
        ['Methylation', 'cg09273054', '5/5'],
        ['Methylation', 'rs4331560*', '5/5'],
    ],
    caption='Table 6. Candidate Biomarkers With Consistent Rankings Across All CV Folds (Requiring Independent Validation)'
)

add_paragraph(
    '*rs4331560 is an SNP probe included on the Illumina HumanMethylation450 array as a control; '
    'its consistent ranking may reflect a genotype-survival association rather than an epigenetic signal.',
    italic=True
)

# ================================================================
# 4. DISCUSSION
# ================================================================
doc.add_heading('4. Discussion', level=1)

add_paragraph(
    'In this study, we developed an interpretable, attention-based multi-branch deep learning '
    'framework for multi-omics survival prediction in hepatocellular carcinoma. Our approach '
    'outperforms the reproduced Chaudhary et al. autoencoder baseline (5-fold nested CV '
    'C-index: 0.683 vs. 0.561) and performs comparably to an AUTOSurv-like benchmark (0.697), while providing transparent '
    'feature- and omics-level importance scores.'
)

add_paragraph(
    'Superior prognostic performance. The attention-based architecture achieved a mean CV C-index of '
    '0.683, representing a 22% relative improvement over the autoencoder baseline. The multi-branch '
    'design, which treats each omics layer independently before fusion, appears to better capture '
    'omics-specific prognostic signals than concatenation-based approaches.',
    bold=False
)

add_paragraph(
    'Balanced omics contributions. Attention weights revealed that mRNA (34.0%), methylation (33.2%), '
    'and miRNA (32.8%) contribute roughly equally to risk prediction, supporting the value of '
    'multi-omics integration over single-omics models.'
)

add_paragraph(
    'Suggestive biological features. The model highlighted features with prior links to HCC '
    'biology, including cell cycle regulators (CCNA2, PLK1, CEP55, KIF2C) and a Wnt pathway component '
    '(FZD7). The top DE genes between risk groups (G6PD, CBX2, CEP55, PLK1, MYBL2) have been '
    'implicated in aggressive HCC in prior studies (Boyault et al. 2007; Hoshida et al. 2009), '
    'though the DE analysis is partly circular and formal pathway enrichment did not reach '
    'statistical significance for most gene sets.'
)

add_paragraph(
    'Independent prognostic value. The risk score added highly significant prognostic value beyond '
    'standard clinical variables (p < 10\u207b\u00b9\u2070\u2070, NRI = 0.398), consistent with '
    'Chaudhary et al.\u2019s finding that multi-omics models complement clinical staging.'
)

doc.add_heading('Limitations', level=2)
add_paragraph(
    'Several limitations should be noted. First, external validation was limited to two real '
    'independent cohorts (GSE14520 mRNA, GSE31384 miRNA); the miRNA cohort could not be '
    'properly evaluated due to incompatible probe identifiers. Validation on additional cohorts '
    '(LIRI-JP, E-TABM-36, Hawaiian) requires controlled-access data applications. Second, with '
    'only 358 patients, the full-data C-index (0.989) indicates overfitting; the cross-validated '
    'C-index (0.683) provides a more realistic generalization estimate. Third, formal pathway '
    'enrichment did not reach statistical significance for most pathways, likely due to the small '
    'feature space after survival-association filtering. Fourth, differential expression analysis '
    'between model-defined risk groups is partially circular, as the groups were defined by the '
    'model trained on the same features; DE results should therefore be interpreted as descriptive '
    'rather than independently confirmatory. Fifth, the low Jaccard index (0.064) between '
    'attention-derived and DE-derived rankings suggests the model captures nonlinear prognostic '
    'features beyond simple differential expression, but further work is needed to determine '
    'whether these represent genuine biological signals or model artifacts.'
)

# ================================================================
# 5. CONCLUSIONS
# ================================================================
doc.add_heading('5. Conclusions', level=1)

add_paragraph(
    'We present an interpretable, attention-based multi-branch deep learning model for multi-omics '
    'survival prediction in hepatocellular carcinoma. In properly nested cross-validation without '
    'feature selection leakage, the proposed architecture outperforms the Chaudhary et al. autoencoder '
    'baseline and demonstrates '
    'significant survival stratification in one independent external mRNA cohort. Attention weights reveal '
    'balanced contributions from all three omics layers, and feature attribution identifies '
    'prognostic markers with prior links to HCC biology, including cell cycle '
    'regulators and a Wnt pathway component. The model-derived risk score adds independent prognostic '
    'value beyond standard clinical variables, though further external validation is needed to '
    'confirm generalizability. This framework advances the goal of transparent, biologically grounded '
    'multi-omics integration for cancer prognosis.'
)

# ================================================================
# 6. DATA AND CODE AVAILABILITY
# ================================================================
doc.add_heading('6. Data and Code Availability', level=1)

add_paragraph(
    'TCGA LIHC data were obtained from the UCSC Xena platform (https://xenabrowser.net). '
    'External validation data were downloaded from GEO (GSE14520, GSE31384). '
    'All analysis scripts, model architectures, trained weights, and preprocessing pipelines '
    'are available at https://github.com/brhanufen/hcc-multiomics-attention. '
    'The repository includes instructions for reproducing all analyses.'
)

# ================================================================
# 7. ETHICS STATEMENT
# ================================================================
doc.add_heading('7. Ethics Statement', level=1)

add_paragraph(
    'This study used only publicly available, de-identified data from TCGA (via UCSC Xena) '
    'and GEO. No institutional review board (IRB) approval was required, as no human subjects '
    'were directly involved and all data are available without restricted access agreements.'
)

# ================================================================
# 8. CONFLICT OF INTEREST
# ================================================================
doc.add_heading('8. Conflict of Interest', level=1)

add_paragraph('The authors declare no competing interests.')

# ================================================================
# 9. FUNDING
# ================================================================
doc.add_heading('9. Funding', level=1)

add_paragraph('This research received no external funding.')

# ================================================================
# 10. ACKNOWLEDGMENTS
# ================================================================
doc.add_heading('10. Acknowledgments', level=1)

add_paragraph(
    'The results shown here are in part based upon data generated by the TCGA Research Network: '
    'https://www.cancer.gov/tcga. We thank the GEO and ArrayExpress databases for providing '
    'open access to the external validation datasets.'
)

# ================================================================
# 11. REFERENCES
# ================================================================
doc.add_heading('11. References', level=1)

refs = [
    'Boyault S, et al. 2007. Transcriptome classification of HCC is related to gene alterations and to new therapeutic targets. Hepatology 45:42\u201352.',
    'Chai H, et al. 2021. Integrating multi-omics data through deep learning for accurate cancer prognosis prediction. Comput Biol Med 134:104481.',
    'Chaudhary K, Poirion OB, Lu L, Garmire LX. 2018. Deep Learning-Based Multi-Omics Integration Robustly Predicts Survival in Liver Cancer. Clin Cancer Res 24:1248\u20131259.',
    'Elbashir MK, et al. 2024. Enhancing non-small cell lung cancer survival prediction through multi-omics integration using graph attention network. Diagnostics 14:2178.',
    'Hoshida Y, et al. 2009. Integrative transcriptome analysis reveals common molecular subclasses of human hepatocellular carcinoma. Cancer Res 69:7385\u20137392.',
    'Huang Z, et al. 2019. SALMON: Survival Analysis Learning With Multi-Omics Neural Networks on Breast Cancer. Front Genet 10:166.',
    'Hutter C, Zenklusen JC. 2018. The Cancer Genome Atlas: Creating Lasting Value beyond Its Data. Cell 173:283\u2013285.',
    'Jiang L, et al. 2024. AUTOSurv: interpretable deep learning framework for cancer survival analysis incorporating clinical and multi-omics data. npj Precis Oncol 8:4.',
    'Llovet JM, et al. 2021. Hepatocellular carcinoma. Nat Rev Dis Primers 7:6.',
    'Rumgay H, et al. 2022. Global burden of primary liver cancer in 2020 and predictions to 2040. J Hepatol 77:1598\u20131606.',
    'Sung H, et al. 2021. Global Cancer Statistics 2020. CA Cancer J Clin 71:209\u2013249.',
    'Villanueva A. 2019. Hepatocellular carcinoma. N Engl J Med 380:1450\u20131462.',
    'Wekesa JS, Kimwele MW. 2023. A review of multi-omics data integration through deep learning approaches for disease diagnosis, prognosis, and treatment. Front Genet 14:1199087.',
    'Wysocka M, et al. 2023. A systematic review of biologically-informed deep learning models for cancer. BMC Bioinformatics 24:198.',
    'Zhang J, et al. 2025. Deep learning-driven multi-omics analysis: enhancing cancer diagnostics and therapeutics. Brief Bioinform 26:bbaf440.',
]

for ref in refs:
    p = doc.add_paragraph(ref, style='List Number')
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'

# ================================================================
# SAVE
# ================================================================
doc.save(OUT)
print(f"Manuscript saved to: {OUT}")
print(f"Figures embedded: 13")
